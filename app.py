import streamlit as st
import google.generativeai as genai
from google.api_core.exceptions import ResourceExhausted, GoogleAPIError
import pdfplumber
import docx
import os
import re
import time
import io
import json
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed
from dotenv import load_dotenv
import plotly.graph_objects as go
import plotly.express as px
import xlsxwriter

load_dotenv()

LOTE_TAMANHO      = 11
PAUSA_LOTE        = 62
ARQUIVO_VAGAS     = os.path.join(os.path.dirname(__file__), "vagas_salvas.json")
ARQUIVO_PERFIS    = os.path.join(os.path.dirname(__file__), "perfis_triagem.json")
ARQUIVO_CRITERIOS = os.path.join(os.path.dirname(__file__), "criterios_salvos.json")
ARQUIVO_CONFIG    = os.path.join(os.path.dirname(__file__), "config_app.json")
ARQUIVO_HISTORICO = os.path.join(os.path.dirname(__file__), "historico_analises.json")

# ── Persistência ───────────────────────────────────────────────────────────────

def _ler(caminho: str) -> dict:
    if os.path.exists(caminho):
        with open(caminho, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}

def _gravar(caminho: str, dados: dict) -> None:
    with open(caminho, "w", encoding="utf-8") as f:
        json.dump(dados, f, ensure_ascii=False, indent=2)

def carregar_vagas() -> dict:   return _ler(ARQUIVO_VAGAS)
def salvar_vagas(d: dict):      _gravar(ARQUIVO_VAGAS, d)
def deletar_vaga(nome: str):
    v = carregar_vagas(); v.pop(nome, None); salvar_vagas(v)

def carregar_perfis() -> dict:    return _ler(ARQUIVO_PERFIS)
def salvar_perfis(d: dict):       _gravar(ARQUIVO_PERFIS, d)
def deletar_perfil(nome: str):
    p = carregar_perfis(); p.pop(nome, None); salvar_perfis(p)

def carregar_criterios() -> dict: return _ler(ARQUIVO_CRITERIOS)
def salvar_criterios(d: dict):    _gravar(ARQUIVO_CRITERIOS, d)
def deletar_criterio(nome: str):
    c = carregar_criterios(); c.pop(nome, None); salvar_criterios(c)

def carregar_config() -> dict: return _ler(ARQUIVO_CONFIG)
def salvar_config(d: dict):    _gravar(ARQUIVO_CONFIG, d)

def carregar_historico() -> list:
    return _ler(ARQUIVO_HISTORICO).get("entradas", [])

def salvar_historico_entrada(entrada: dict):
    dados = _ler(ARQUIVO_HISTORICO)
    entradas = dados.get("entradas", [])
    entradas.insert(0, entrada)
    _gravar(ARQUIVO_HISTORICO, {"entradas": entradas[:20]})

# ── Setup da página ────────────────────────────────────────────────────────────

st.set_page_config(
    page_title="Avaliador de Currículos",
    page_icon="📋",
    layout="wide",
)

if "score_corte" not in st.session_state:
    cfg = carregar_config()
    st.session_state["score_corte"] = cfg.get("score_corte", 75)

# ── Sidebar: Histórico de Triagens ─────────────────────────────────────────────

with st.sidebar:
    st.header("🕓 Histórico de Triagens")
    historico = carregar_historico()
    if not historico:
        st.caption("Nenhuma triagem salva ainda.")
    else:
        for i, entrada in enumerate(historico):
            with st.expander(f"📅 {entrada['timestamp']} — {entrada['total']} CVs"):
                st.caption(entrada["vaga_preview"])
                st.caption(f"✅ {entrada['aprovados']} aprovados · Score mínimo: {entrada['score_corte']}")
                if st.button("📂 Carregar esta triagem", key=f"hist_{i}", use_container_width=True):
                    st.session_state["resultados"]        = entrada["resultados"]
                    st.session_state["descricao_vaga"]    = entrada["descricao_vaga"]
                    st.session_state["score_corte_usado"] = entrada["score_corte"]
                    st.session_state["criterios_usados"]  = entrada.get("criterios_usados", {})
                    st.rerun()

    st.divider()
    if st.session_state.get("resultados"):
        if st.button("🗑️ Limpar Resultados Atuais", use_container_width=True):
            for k in ["resultados", "descricao_vaga", "score_corte_usado", "ranking_base",
                      "comparacao", "criterios_usados"]:
                st.session_state.pop(k, None)
            for k in list(st.session_state.keys()):
                if k.startswith("roteiro_"):
                    del st.session_state[k]
            st.rerun()

st.title("📋 Avaliador de Currículos")
st.caption("Analise candidatos com base na vaga e nos critérios esperados")

# ── Funções de extração de texto ───────────────────────────────────────────────

def extrair_texto_pdf(arquivo) -> str:
    texto = []
    with pdfplumber.open(arquivo) as pdf:
        for pagina in pdf.pages:
            conteudo = pagina.extract_text()
            if conteudo:
                texto.append(conteudo)
    return "\n".join(texto)

def extrair_texto_docx(arquivo) -> str:
    arquivo.seek(0)
    doc = docx.Document(arquivo)
    linhas = []
    for p in doc.paragraphs:
        if p.text.strip():
            linhas.append(p.text.strip())
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for p in celula.paragraphs:
                    if p.text.strip():
                        linhas.append(p.text.strip())
    return "\n".join(linhas)

def extrair_texto(arquivo) -> str:
    arquivo.seek(0)
    nome = arquivo.name.lower()
    if nome.endswith(".pdf"):
        return extrair_texto_pdf(arquivo)
    elif nome.endswith(".docx") or nome.endswith(".doc"):
        return extrair_texto_docx(arquivo)
    return arquivo.read().decode("utf-8", errors="ignore")

def extrair_score(texto: str) -> int:
    match = re.search(r'\b(\d{1,3})\s*(?:/\s*100|pontos?)?', texto[:500])
    if match:
        valor = int(match.group(1))
        if 0 <= valor <= 100:
            return valor
    return 0

def extrair_veredicto(texto: str) -> str:
    match = re.search(r'VEREDICTO\s*:\s*(Aprovado|Reprovado)', texto, re.IGNORECASE)
    if match:
        return match.group(1).capitalize()
    return "Indefinido"

def extrair_scores_criterios(texto: str, criterios: list) -> dict:
    scores = {}
    for crit in criterios:
        match = re.search(
            rf'\*\*{re.escape(crit)}\*\*\s*:\s*NOTA:\s*(\d+)',
            texto, re.IGNORECASE
        )
        scores[crit] = int(match.group(1)) if match else 0
    return scores

def medalha(posicao: int) -> str:
    return {1: "🥇", 2: "🥈", 3: "🥉"}.get(posicao, f"{posicao}º")

def badge_veredicto(veredicto: str) -> str:
    if veredicto == "Aprovado":
        return "🟢 Aprovado"
    if veredicto == "Reprovado":
        return "🔴 Reprovado"
    return "⚪ Indefinido"

def extrair_secao(texto: str, titulo: str) -> str:
    pattern = rf"###?\s*[^\n]*{re.escape(titulo)}[^\n]*\n(.*?)(?=###|\Z)"
    match = re.search(pattern, texto, re.DOTALL | re.IGNORECASE)
    if match:
        return match.group(1).strip()[:500]
    return ""

def gerar_excel(ranking: list, resultados: dict) -> bytes:
    output = io.BytesIO()
    wb = xlsxwriter.Workbook(output, {"in_memory": True})
    ws = wb.add_worksheet("Ranking")

    fmt_header    = wb.add_format({"bold": True, "bg_color": "#1a1a2e", "font_color": "#ffffff", "border": 1, "align": "center", "valign": "vcenter"})
    fmt_aprovado  = wb.add_format({"bg_color": "#d4edda", "font_color": "#155724", "bold": True, "border": 1, "align": "center"})
    fmt_reprovado = wb.add_format({"bg_color": "#f8d7da", "font_color": "#721c24", "bold": True, "border": 1, "align": "center"})
    fmt_normal    = wb.add_format({"border": 1, "valign": "top", "text_wrap": True})
    fmt_center    = wb.add_format({"border": 1, "align": "center", "valign": "vcenter", "bold": True})
    fmt_score_alto  = wb.add_format({"border": 1, "align": "center", "valign": "vcenter", "bold": True, "font_color": "#155724"})
    fmt_score_med   = wb.add_format({"border": 1, "align": "center", "valign": "vcenter", "bold": True, "font_color": "#856404"})
    fmt_score_baixo = wb.add_format({"border": 1, "align": "center", "valign": "vcenter", "bold": True, "font_color": "#721c24"})

    headers = ["Posição", "Candidato", "Score", "Veredicto", "Pontos Fortes", "Lacunas", "Recomendação"]
    widths  = [10, 25, 10, 14, 45, 45, 45]

    for col, (h, w) in enumerate(zip(headers, widths)):
        ws.set_column(col, col, w)
        ws.write(0, col, h, fmt_header)
    ws.set_row(0, 20)

    for i, (nome, score, veredicto) in enumerate(ranking):
        row = i + 1
        ws.set_row(row, 80)
        nome_curto = nome.rsplit(".", 1)[0]
        analise = resultados[nome]["analise"]
        fortes  = extrair_secao(analise, "Pontos Fortes")
        lacunas = extrair_secao(analise, "Lacunas")
        rec     = extrair_secao(analise, "Recomendação Final")

        fmt_score = fmt_score_alto if score >= 80 else (fmt_score_med if score >= 60 else fmt_score_baixo)
        fmt_verd  = fmt_aprovado if veredicto == "Aprovado" else fmt_reprovado

        ws.write(row, 0, i + 1, fmt_center)
        ws.write(row, 1, nome_curto, fmt_normal)
        ws.write(row, 2, f"{score}/100", fmt_score)
        ws.write(row, 3, veredicto, fmt_verd)
        ws.write(row, 4, fortes, fmt_normal)
        ws.write(row, 5, lacunas, fmt_normal)
        ws.write(row, 6, rec, fmt_normal)

    wb.close()
    return output.getvalue()

# ── API helpers ────────────────────────────────────────────────────────────────

def testar_api_key() -> tuple[bool, str]:
    api_key = os.getenv("GOOGLE_API_KEY", "")
    if not api_key or api_key == "sua_chave_aqui":
        return False, "⚠️ API Key não configurada. Verifique o arquivo `.env`."
    try:
        genai.configure(api_key=api_key)
        modelo = genai.GenerativeModel("gemini-3.6-flash")
        modelo.generate_content("Responda apenas: OK")
        return True, ""
    except ResourceExhausted:
        return False, "⚠️ Rate limit atingido. Aguarde 1 minuto e tente novamente."
    except GoogleAPIError as e:
        return False, f"⚠️ Erro na API do Google: {str(e)}"
    except Exception as e:
        return False, f"⚠️ Não foi possível conectar à API: {str(e)}"


def analisar_curriculo(nome_arquivo: str, texto_cv: str, descricao_vaga: str, expectativas: str, criterios: dict = None, score_corte: int = 75) -> tuple[str, str, str]:
    api_key = os.getenv("GOOGLE_API_KEY", "")
    if not api_key or api_key == "sua_chave_aqui":
        msg = "⚠️ **API Key não configurada.** Edite o arquivo `.env` e insira sua chave do Google AI Studio."
        return nome_arquivo, msg, msg

    genai.configure(api_key=api_key)
    modelo = genai.GenerativeModel("gemini-3.6-flash")

    if not criterios:
        criterios = {"Habilidades Técnicas": 40, "Experiência Relevante": 25,
                     "Formação Acadêmica": 10, "Idiomas": 10, "Soft Skills": 10, "Fit Cultural": 5}

    lista_criterios = "\n".join(f"- {c} (peso: {p}%)" for c, p in criterios.items())

    prompt_analise = f"""Você é um especialista em recrutamento e seleção. Analise o currículo abaixo com base na vaga e nas expectativas fornecidas.

---
## CURRÍCULO DO CANDIDATO
{texto_cv}

---
## DESCRIÇÃO DA VAGA E REQUISITOS
{descricao_vaga}

---
## TRIAGEM DE PERFIL
{expectativas}

---

Elabore uma análise completa em português com as seguintes seções:

### 📊 Avaliação por Critério
Para cada critério abaixo, atribua uma nota de 0 a 100 e uma justificativa curta:
{lista_criterios}

Use o formato exato para cada critério:
**[Nome do Critério]**: NOTA: XX — justificativa

### 🎯 Score Geral
Calcule o score ponderado usando os pesos acima e apresente no formato exato:
Score: XX/100
Depois inclua a classificação: **Altamente Recomendado** / **Recomendado** / **Recomendado com Ressalvas** / **Não Recomendado**

### ✅ Pontos Fortes
Liste os principais pontos positivos do candidato em relação à vaga (bullets).

### ⚠️ Lacunas Identificadas
Liste requisitos da vaga que o candidato não atende ou atende parcialmente (bullets).

### 💬 Recomendação Final
Um parágrafo com a recomendação objetiva sobre avançar ou não com o candidato.

### ✅ Veredicto
Com base em todos os critérios acima, emita o veredicto final obrigatoriamente no formato exato:
VEREDICTO: Aprovado
ou
VEREDICTO: Reprovado

Critério: candidatos com score abaixo de {score_corte} ou que não atendem requisitos obrigatórios devem ser Reprovados.
"""

    prompt_unico = prompt_analise + f"""

---
---

## PARTE 2 — PERGUNTAS DE ENTREVISTA

Agora, como recrutador sênior, gere perguntas de entrevista personalizadas para este candidato:

### 🔧 Perguntas Técnicas (5 perguntas)
Perguntas aprofundadas sobre as competências técnicas exigidas, explorando especialmente as lacunas identificadas.

### 🧠 Perguntas Comportamentais (4 perguntas)
Perguntas no formato STAR (Situação, Tarefa, Ação, Resultado) para avaliar soft skills e fit cultural.

### 🔍 Perguntas de Aprofundamento (3 perguntas)
Perguntas específicas para esclarecer pontos do currículo que ficaram vagos ou que merecem mais detalhes.

---
SEPARE as duas partes com exatamente esta linha:
===PERGUNTAS===
"""

    try:
        resposta = modelo.generate_content(prompt_unico)
        texto_completo = resposta.text
        if "===PERGUNTAS===" in texto_completo:
            partes    = texto_completo.split("===PERGUNTAS===", 1)
            analise   = partes[0].strip()
            perguntas = partes[1].strip()
        else:
            analise   = texto_completo
            perguntas = "_Perguntas não foram geradas nesta análise._"
        return nome_arquivo, analise, perguntas
    except ResourceExhausted:
        msg = (
            "⚠️ **Limite de uso da API atingido (Rate Limit)**\n\n"
            "O plano gratuito do Gemini permite até **15 chamadas por minuto**.\n\n"
            "**O que fazer:** Aguarde 1 minuto e tente novamente."
        )
        return nome_arquivo, msg, msg
    except GoogleAPIError as e:
        msg = f"⚠️ **Erro na API do Google:** {str(e)}"
        return nome_arquivo, msg, msg
    except Exception as e:
        msg = f"⚠️ **Erro inesperado ao processar este currículo:** {str(e)}"
        return nome_arquivo, msg, msg


def comparar_candidatos(selecionados: list[str], resultados: dict, descricao_vaga: str) -> str:
    api_key = os.getenv("GOOGLE_API_KEY", "")
    genai.configure(api_key=api_key)
    modelo = genai.GenerativeModel("gemini-3.6-flash")

    blocos = ""
    for nome in selecionados:
        nome_curto = nome.rsplit(".", 1)[0]
        blocos += f"\n\n### Candidato: {nome_curto}\n{resultados[nome]['analise']}"

    prompt = f"""Você é um especialista em recrutamento. Compare os candidatos abaixo para a vaga descrita e produza uma análise comparativa em português.

## VAGA
{descricao_vaga}

## ANÁLISES DOS CANDIDATOS
{blocos}

---

Produza a comparação com as seguintes seções:

### 🏆 Ranking Comparativo
Ordene os candidatos do mais adequado ao menos adequado com justificativa em 1 linha cada.

### 📊 Tabela Comparativa
Crie uma tabela markdown comparando os candidatos nos principais critérios da vaga.
Use o formato:
| Critério | Candidato A | Candidato B | ... |
|---|---|---|---|

### ✅ Pontos Fortes de Cada Um
Para cada candidato, 2-3 diferenciais em relação aos demais.

### ⚠️ Principal Risco de Cada Um
Para cada candidato, o maior ponto de atenção.

### 💡 Recomendação Final
Qual candidato avançar e por quê, em um parágrafo objetivo.
"""

    try:
        resposta = modelo.generate_content(prompt)
        return resposta.text
    except ResourceExhausted:
        return "⚠️ **Limite da API atingido.** Aguarde 1 minuto e tente novamente."
    except Exception as e:
        return f"⚠️ **Erro ao comparar:** {str(e)}"


def gerar_roteiro_entrevista(texto_cv: str, analise: str, descricao_vaga: str, expectativas: str) -> str:
    api_key = os.getenv("GOOGLE_API_KEY", "")
    genai.configure(api_key=api_key)
    modelo = genai.GenerativeModel("gemini-3.6-flash")

    prompt = f"""Você é um especialista em recrutamento e seleção. Com base no currículo, na análise já realizada e na vaga, crie um roteiro completo e estruturado para conduzir a entrevista com este candidato.

---
## CURRÍCULO DO CANDIDATO
{texto_cv}

---
## ANÁLISE JÁ REALIZADA
{analise}

---
## DESCRIÇÃO DA VAGA
{descricao_vaga}

---
## TRIAGEM DE PERFIL
{expectativas}

---

Gere o roteiro em português com as seguintes seções:

### 👋 1. Abertura (5 min)
Perguntas de quebra-gelo e rapport. Objetivo: deixar o candidato confortável e iniciar a conversa.
Inclua 2-3 perguntas sugeridas.

### 🔍 2. Validação de Experiências e Desafios Profissionais (15 min)
Com base nos desafios e realizações mencionados no CV, elabore perguntas para validar a veracidade e profundidade de cada experiência relevante.
Para cada experiência/projeto mencionado no CV, inclua:
- A experiência citada no CV
- 2 perguntas para validar e aprofundar

### 🎯 3. Validação de Senioridade (10 min)
Perguntas para confirmar se a senioridade declarada no CV condiz com o nível real de conhecimento e autonomia.
Inclua perguntas que diferenciem claramente um profissional júnior, pleno e sênior — adequadas ao perfil esperado.

### 🔧 4. Competências Técnicas Obrigatórias (15 min)
Para cada requisito técnico obrigatório da vaga, elabore 1-2 perguntas práticas ou cenários para avaliar o domínio real.

### 🧠 5. Competências Comportamentais e Fit Cultural (10 min)
Perguntas no formato STAR para avaliar soft skills e aderência ao perfil esperado.

### 🚀 6. Motivação e Expectativas (5 min)
Perguntas sobre motivação, expectativas de carreira e alinhamento com a vaga.

### ✅ 7. Encerramento (5 min)
Perguntas finais e espaço para o candidato tirar dúvidas.
Inclua orientação sobre os próximos passos a comunicar ao candidato.

### 📋 8. Guia de Avaliação
Para cada seção do roteiro, indique:
- O que avaliar
- Sinais positivos (o que o bom candidato deve demonstrar)
- Sinais de alerta (respostas que indicam riscos)
"""

    try:
        resposta = modelo.generate_content(prompt)
        return resposta.text
    except ResourceExhausted:
        return "⚠️ **Limite da API atingido.** Aguarde 1 minuto e tente novamente."
    except Exception as e:
        return f"⚠️ **Erro ao gerar roteiro:** {str(e)}"


# ── Interface ──────────────────────────────────────────────────────────────────

col_esq, col_dir = st.columns(2)

with col_esq:
    st.subheader("📄 Currículos dos Candidatos")
    arquivos = st.file_uploader(
        "Faça o upload dos currículos",
        type=["pdf", "docx", "doc", "txt"],
        accept_multiple_files=True,
        help="Formatos aceitos: PDF, DOCX, DOC, TXT — múltiplos arquivos permitidos",
    )
    if arquivos:
        st.success(f"{len(arquivos)} arquivo(s) carregado(s):")
        for a in arquivos:
            st.caption(f"• {a.name}")

with col_dir:
    st.subheader("💼 Descrição da Vaga & Requisitos")

    vagas = carregar_vagas()
    if vagas:
        v_sel = st.selectbox("📂 Carregar vaga salva", ["— selecione —"] + list(vagas.keys()), key="sel_vaga")
        cv1, cv2 = st.columns(2)
        with cv1:
            if st.button("Carregar", key="btn_carregar_vaga", use_container_width=True) and v_sel != "— selecione —":
                st.session_state["desc_vaga"] = vagas[v_sel]
                st.rerun()
        with cv2:
            if st.button("🗑️ Excluir", key="btn_del_vaga", use_container_width=True) and v_sel != "— selecione —":
                deletar_vaga(v_sel)
                st.rerun()

    descricao_vaga = st.text_area(
        "Cole aqui a descrição da vaga, responsabilidades e requisitos",
        height=140,
        placeholder="Ex: Desenvolvedor Full Stack Sênior\n\nRequisitos:\n• 5+ anos com React e Node.js\n• Experiência com AWS",
        key="desc_vaga",
    )
    with st.expander("💾 Salvar esta descrição de vaga"):
        nome_vaga = st.text_input("Nome (ex: Analista de Dados — Jul/2026)", key="nome_nova_vaga")
        if st.button("💾 Salvar Vaga", use_container_width=True, key="salvar_vaga"):
            if not nome_vaga.strip():
                st.warning("Digite um nome.")
            elif not descricao_vaga.strip():
                st.warning("Preencha a descrição antes de salvar.")
            else:
                v = carregar_vagas()
                v[nome_vaga.strip()] = descricao_vaga
                salvar_vagas(v)
                st.success(f'✅ Vaga "{nome_vaga}" salva!')
                st.rerun()

    st.subheader("🎯 Triagem de Perfil")

    perfis = carregar_perfis()
    if perfis:
        p_sel = st.selectbox("📂 Carregar perfil salvo", ["— selecione —"] + list(perfis.keys()), key="sel_perfil")
        cp1, cp2 = st.columns(2)
        with cp1:
            if st.button("Carregar", key="btn_carregar_perfil", use_container_width=True) and p_sel != "— selecione —":
                st.session_state["triagem"] = perfis[p_sel]
                st.rerun()
        with cp2:
            if st.button("🗑️ Excluir", key="btn_del_perfil", use_container_width=True) and p_sel != "— selecione —":
                deletar_perfil(p_sel)
                st.rerun()

    expectativas = st.text_area(
        "Descreva o perfil comportamental, soft skills e fit cultural esperados",
        height=110,
        placeholder="Ex: Perfil proativo, acostumado com squads ágeis, capacidade de liderança técnica...",
        key="triagem",
    )
    with st.expander("💾 Salvar este perfil de triagem"):
        nome_perfil = st.text_input("Nome (ex: Perfil Sênior Técnico)", key="nome_novo_perfil")
        if st.button("💾 Salvar Perfil", use_container_width=True, key="salvar_perfil"):
            if not nome_perfil.strip():
                st.warning("Digite um nome.")
            elif not expectativas.strip():
                st.warning("Preencha a triagem antes de salvar.")
            else:
                p = carregar_perfis()
                p[nome_perfil.strip()] = expectativas
                salvar_perfis(p)
                st.success(f'✅ Perfil "{nome_perfil}" salvo!')
                st.rerun()

st.divider()

# ── Critérios de Avaliação ─────────────────────────────────────────────────────
CRITERIOS_PADRAO = {
    "Habilidades Técnicas":  40,
    "Experiência Relevante": 25,
    "Formação Acadêmica":    10,
    "Idiomas":               10,
    "Soft Skills":           10,
    "Fit Cultural":           5,
}

with st.expander("⚖️ Configurar Critérios e Pesos de Avaliação", expanded=False):
    st.caption("Selecione os critérios que importam para esta vaga e ajuste os pesos. O total deve somar 100%.")

    configs_salvas = carregar_criterios()
    if configs_salvas:
        cc1, cc2, cc3 = st.columns([3, 1, 1])
        with cc1:
            config_sel = st.selectbox("📂 Carregar configuração salva", ["— selecione —"] + list(configs_salvas.keys()), key="sel_config_crit")
        with cc2:
            st.write("")
            st.write("")
            if st.button("Carregar", key="btn_carregar_crit", use_container_width=True) and config_sel != "— selecione —":
                cfg = configs_salvas[config_sel]
                st.session_state["criterios_sel"] = list(cfg.keys())
                for c, p in cfg.items():
                    st.session_state[f"peso_{c}"] = p
                st.rerun()
        with cc3:
            st.write("")
            st.write("")
            if st.button("🗑️ Excluir", key="btn_del_crit", use_container_width=True) and config_sel != "— selecione —":
                deletar_criterio(config_sel)
                st.rerun()

    st.divider()

    criterios_disponiveis = list(CRITERIOS_PADRAO.keys()) + ["Liderança", "Gestão de Projetos", "Vendas / Relacionamento"]
    selecionados = st.multiselect(
        "Critérios ativos",
        options=criterios_disponiveis,
        default=st.session_state.get("criterios_sel", list(CRITERIOS_PADRAO.keys())),
        key="criterios_sel",
    )

    pesos = {}
    if selecionados:
        st.markdown("**Defina o peso de cada critério (total deve ser 100%):**")
        cols_crit = st.columns(min(len(selecionados), 3))
        for i, crit in enumerate(selecionados):
            with cols_crit[i % 3]:
                peso_default = st.session_state.get(f"peso_{crit}", CRITERIOS_PADRAO.get(crit, 10))
                pesos[crit] = st.number_input(
                    f"{crit} (%)",
                    min_value=0, max_value=100,
                    value=peso_default,
                    step=5,
                    key=f"peso_{crit}",
                )

        total_pesos = sum(pesos.values())
        if total_pesos == 100:
            st.success(f"✅ Total: {total_pesos}% — pronto para análise!")
        else:
            diff = 100 - total_pesos
            st.warning(f"⚠️ Total: {total_pesos}% — {'falta' if diff > 0 else 'excede'} {abs(diff)}% para chegar em 100%.")

        st.divider()
        sc1, sc2 = st.columns([3, 1])
        with sc1:
            nome_config = st.text_input("Nome da configuração (ex: Vaga Técnica Sênior)", key="nome_nova_config")
        with sc2:
            st.write("")
            st.write("")
            if st.button("💾 Salvar", key="salvar_config_crit", use_container_width=True):
                if not nome_config.strip():
                    st.warning("Digite um nome.")
                elif total_pesos != 100:
                    st.warning("O total precisa ser 100% para salvar.")
                else:
                    configs = carregar_criterios()
                    configs[nome_config.strip()] = pesos
                    salvar_criterios(configs)
                    st.success(f'✅ Configuração "{nome_config}" salva!')
                    st.rerun()
    else:
        st.warning("Selecione pelo menos um critério.")

criterios_para_analise = pesos if pesos and sum(pesos.values()) == 100 else CRITERIOS_PADRAO

col_score, col_btn = st.columns([1, 3])
with col_score:
    score_corte = st.number_input(
        "Score mínimo para aprovação",
        min_value=0, max_value=100,
        step=5,
        help="Candidatos com score abaixo deste valor serão automaticamente Reprovados",
        key="score_corte",
    )
with col_btn:
    analisar = st.button("🤖 Analisar Currículos", type="primary", use_container_width=True)

if analisar:
    if not arquivos:
        st.warning("Faça o upload de pelo menos um currículo antes de analisar.")
    elif not descricao_vaga.strip():
        st.warning("Preencha a descrição da vaga.")
    elif not expectativas.strip():
        st.warning("Preencha a triagem de perfil.")
    else:
        # ── Validar API key antes de iniciar ──────────────────────────────────
        with st.spinner("🔌 Verificando conexão com a API Gemini..."):
            api_ok, api_msg = testar_api_key()
        if not api_ok:
            st.error(api_msg)
            st.stop()

        salvar_config({"score_corte": score_corte})

        candidatos = [(a.name, extrair_texto(a)) for a in arquivos]
        total      = len(candidatos)
        lotes      = [candidatos[i:i + LOTE_TAMANHO] for i in range(0, total, LOTE_TAMANHO)]
        num_lotes  = len(lotes)

        st.info(
            f"📦 **{total} CVs** divididos em **{num_lotes} lote(s)** de até {LOTE_TAMANHO} "
            f"para respeitar o limite da API."
        )

        progress   = st.progress(0, text="Iniciando...")
        status_txt = st.empty()
        resultados = {}
        erros      = []
        concluidos = 0

        PAUSA_ENTRE_CVS = 4

        for idx_lote, lote in enumerate(lotes, start=1):
            status_txt.markdown(f"🔄 **Lote {idx_lote}/{num_lotes}** — analisando {len(lote)} currículo(s)...")

            cvs_rate_limit = []

            for idx_cv, (nome, texto) in enumerate(lote):
                status_txt.markdown(
                    f"🔄 **Lote {idx_lote}/{num_lotes}** — analisando: {nome} ({idx_cv+1}/{len(lote)})"
                )
                nome_arquivo, analise, perguntas = analisar_curriculo(
                    nome, texto, descricao_vaga, expectativas, criterios_para_analise, score_corte
                )
                if analise.startswith("⚠️ **Limite"):
                    cvs_rate_limit.append((nome, texto))
                else:
                    resultados[nome_arquivo] = {"analise": analise, "perguntas": perguntas, "texto_cv": texto}
                    concluidos += 1
                    progress.progress(concluidos / total, text=f"Concluído: {nome_arquivo} ({concluidos}/{total})")

                if idx_cv < len(lote) - 1:
                    time.sleep(PAUSA_ENTRE_CVS)

            if cvs_rate_limit:
                for seg in range(65, 0, -1):
                    status_txt.markdown(
                        f"⏳ Rate limit atingido. Aguardando **{seg}s** para retentar "
                        f"{len(cvs_rate_limit)} CV(s)..."
                    )
                    time.sleep(1)
                for nome, texto in cvs_rate_limit:
                    status_txt.markdown(f"🔄 Retentando: {nome}...")
                    nome_arquivo, analise, perguntas = analisar_curriculo(
                        nome, texto, descricao_vaga, expectativas, criterios_para_analise, score_corte
                    )
                    resultados[nome_arquivo] = {"analise": analise, "perguntas": perguntas, "texto_cv": texto}
                    concluidos += 1
                    progress.progress(concluidos / total, text=f"Concluído: {nome_arquivo} ({concluidos}/{total})")
                    time.sleep(PAUSA_ENTRE_CVS)

            if idx_lote < num_lotes:
                for seg in range(PAUSA_LOTE, 0, -1):
                    status_txt.markdown(
                        f"⏳ Lote {idx_lote}/{num_lotes} concluído. "
                        f"Aguardando **{seg}s** antes do próximo lote..."
                    )
                    time.sleep(1)

        status_txt.empty()
        progress.empty()

        if erros:
            st.warning(f"⚠️ {len(erros)} currículo(s) com erro: {', '.join(erros)}.")

        # aplica score mínimo client-side
        for nome_r, dados_r in resultados.items():
            if extrair_score(dados_r["analise"]) < score_corte:
                dados_r["analise"] = re.sub(
                    r'(Veredicto|Veredito|Resultado)\s*:\s*(Aprovado)',
                    r'\1: Reprovado',
                    dados_r["analise"],
                    flags=re.IGNORECASE,
                )

        st.session_state["resultados"]        = resultados
        st.session_state["descricao_vaga"]    = descricao_vaga
        st.session_state["score_corte_usado"] = score_corte
        st.session_state["criterios_usados"]  = criterios_para_analise

        # salvar no histórico
        aprovados_hist = sum(
            1 for d in resultados.values()
            if extrair_veredicto(d["analise"]) == "Aprovado"
        )
        salvar_historico_entrada({
            "timestamp":      datetime.now().strftime("%d/%m/%Y %H:%M"),
            "vaga_preview":   (descricao_vaga[:80] + "...") if len(descricao_vaga) > 80 else descricao_vaga,
            "total":          len(resultados),
            "aprovados":      aprovados_hist,
            "score_corte":    score_corte,
            "criterios_usados": criterios_para_analise,
            "resultados":     resultados,
            "descricao_vaga": descricao_vaga,
        })

# ── Exibe resultados ───────────────────────────────────────────────────────────
if st.session_state.get("resultados"):
    resultados     = st.session_state["resultados"]
    descricao_vaga = st.session_state.get("descricao_vaga", "")

    _score_corte = st.session_state.get("score_corte_usado", st.session_state.get("score_corte", 75))
    ranking_base = []
    for nome, dados in resultados.items():
        score     = extrair_score(dados["analise"])
        veredicto = extrair_veredicto(dados["analise"])
        if score < _score_corte:
            veredicto = "Reprovado"
        ranking_base.append((nome, score, veredicto))
    ranking_base.sort(key=lambda x: x[1], reverse=True)
    st.session_state["ranking_base"] = ranking_base

    # ── Dashboard ─────────────────────────────────────────────────────────────
    st.subheader("📈 Dashboard")

    total_c    = len(ranking_base)
    aprovados  = sum(1 for _, _, v in ranking_base if v == "Aprovado")
    reprovados = total_c - aprovados
    scores     = [s for _, s, _ in ranking_base]
    score_med  = round(sum(scores) / total_c) if total_c else 0

    m1, m2, m3, m4 = st.columns(4)
    m1.metric("Total de CVs",  total_c)
    m2.metric("✅ Aprovados",   aprovados)
    m3.metric("❌ Reprovados",  reprovados)
    m4.metric("Score Médio",   f"{score_med}/100")

    c_pizza, c_barra = st.columns(2)

    with c_pizza:
        fig_pie = go.Figure(go.Pie(
            labels=["Aprovados", "Reprovados"],
            values=[aprovados, reprovados],
            marker_colors=["#69db7c", "#ff6b6b"],
            hole=0.4,
        ))
        fig_pie.update_layout(
            title="Aprovados vs Reprovados",
            margin=dict(t=40, b=10, l=10, r=10),
            height=280,
            showlegend=True,
            paper_bgcolor="rgba(0,0,0,0)",
            font_color="#e8e8f0",
        )
        st.plotly_chart(fig_pie, use_container_width=True)

    with c_barra:
        nomes_curtos = [n.rsplit(".", 1)[0][:15] for n, _, _ in ranking_base]
        cores = ["#69db7c" if v == "Aprovado" else "#ff6b6b" for _, _, v in ranking_base]
        fig_bar = go.Figure(go.Bar(
            x=nomes_curtos,
            y=[s for _, s, _ in ranking_base],
            marker_color=cores,
            text=[f"{s}" for _, s, _ in ranking_base],
            textposition="outside",
        ))
        fig_bar.update_layout(
            title="Score por Candidato",
            yaxis=dict(range=[0, 110]),
            margin=dict(t=40, b=10, l=10, r=10),
            height=280,
            paper_bgcolor="rgba(0,0,0,0)",
            plot_bgcolor="rgba(0,0,0,0)",
            font_color="#e8e8f0",
        )
        st.plotly_chart(fig_bar, use_container_width=True)

    # Distribuição de scores por faixa
    faixas_labels = ["0–25", "26–50", "51–75", "76–100"]
    faixas_vals   = [0, 0, 0, 0]
    for _, s, _ in ranking_base:
        if s <= 25:   faixas_vals[0] += 1
        elif s <= 50: faixas_vals[1] += 1
        elif s <= 75: faixas_vals[2] += 1
        else:         faixas_vals[3] += 1

    fig_hist = go.Figure(go.Bar(
        x=faixas_labels,
        y=faixas_vals,
        marker_color=["#ff6b6b", "#ffa94d", "#ffd43b", "#69db7c"],
        text=faixas_vals,
        textposition="outside",
    ))
    fig_hist.update_layout(
        title="Distribuição de Scores por Faixa",
        yaxis=dict(range=[0, max(faixas_vals) + 1.5]),
        xaxis_title="Faixa de Score",
        yaxis_title="Candidatos",
        margin=dict(t=40, b=10, l=10, r=10),
        height=260,
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(0,0,0,0)",
        font_color="#e8e8f0",
    )
    st.plotly_chart(fig_hist, use_container_width=True)

    st.divider()

    # ── Filtros + Ranking ──────────────────────────────────────────────────────
    st.subheader("🏆 Ranking de Candidatos")

    criterios_usados = st.session_state.get("criterios_usados", {})
    opcoes_ordenacao = ["Score Total"] + list(criterios_usados.keys())

    f1, f2, f3, f4 = st.columns([2, 2, 2, 1])
    with f1:
        score_min = st.slider("Score mínimo", 0, 100, 0, step=5, key="filtro_score")
    with f2:
        filtro_verd = st.radio("Veredicto", ["Todos", "Aprovados", "Reprovados"], horizontal=True, key="filtro_verd")
    with f3:
        ordenar_por = st.selectbox("Ordenar por", opcoes_ordenacao, key="ordenar_por")
    with f4:
        excel_bytes = gerar_excel(ranking_base, resultados)
        st.download_button(
            label="⬇️ Excel",
            data=excel_bytes,
            file_name="ranking_candidatos.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

    # ordenação por critério específico
    if ordenar_por != "Score Total" and criterios_usados:
        ranking_display = sorted(
            ranking_base,
            key=lambda x: extrair_scores_criterios(resultados[x[0]]["analise"], [ordenar_por]).get(ordenar_por, 0),
            reverse=True,
        )
    else:
        ranking_display = ranking_base

    ranking_filtrado = [
        (n, s, v) for n, s, v in ranking_display
        if s >= score_min and (
            filtro_verd == "Todos"
            or (filtro_verd == "Aprovados"  and v == "Aprovado")
            or (filtro_verd == "Reprovados" and v == "Reprovado")
        )
    ]

    if not ranking_filtrado:
        st.info("Nenhum candidato corresponde aos filtros aplicados.")
    else:
        cols = st.columns(min(len(ranking_filtrado), 5))
        for i, (nome, score, veredicto) in enumerate(ranking_filtrado):
            col_idx    = i % len(cols)
            nome_curto = nome.rsplit(".", 1)[0]
            pos_real   = ranking_base.index((nome, score, veredicto)) + 1
            with cols[col_idx]:
                label_extra = ""
                if ordenar_por != "Score Total" and criterios_usados:
                    nota_crit = extrair_scores_criterios(resultados[nome]["analise"], [ordenar_por]).get(ordenar_por, 0)
                    label_extra = f"{ordenar_por[:12]}: {nota_crit}"
                st.metric(
                    label=f"{medalha(pos_real)} {nome_curto}",
                    value=f"{score}/100",
                    delta=label_extra if label_extra else badge_veredicto(veredicto),
                    delta_color="normal" if veredicto == "Aprovado" else "inverse",
                )

    st.divider()

    # ── Comparar Candidatos ────────────────────────────────────────────────────
    st.subheader("🔍 Comparar Candidatos")
    nomes_disponiveis = [nome.rsplit(".", 1)[0] for nome in resultados.keys()]
    nomes_para_chave  = {nome.rsplit(".", 1)[0]: nome for nome in resultados.keys()}

    selecionados_curtos = st.multiselect(
        "Selecione 2 a 4 candidatos para comparar",
        options=nomes_disponiveis,
        max_selections=4,
        placeholder="Escolha os candidatos...",
    )

    if len(selecionados_curtos) >= 2:
        if st.button("🔍 Comparar Selecionados", use_container_width=True):
            selecionados_chaves = [nomes_para_chave[n] for n in selecionados_curtos]
            with st.spinner(f"Comparando {len(selecionados_curtos)} candidatos..."):
                comparacao = comparar_candidatos(selecionados_chaves, resultados, descricao_vaga)
            st.session_state["comparacao"] = comparacao
    elif len(selecionados_curtos) == 1:
        st.caption("Selecione pelo menos 2 candidatos para comparar.")

    if "comparacao" in st.session_state:
        st.markdown(st.session_state["comparacao"])
        st.download_button(
            label="⬇️ Baixar Comparação (TXT)",
            data=st.session_state["comparacao"],
            file_name="comparacao_candidatos.txt",
            mime="text/plain",
            key="dl_comparacao",
        )

    st.divider()

    # ── Abas por candidato ─────────────────────────────────────────────────────
    st.subheader(f"📋 Análise Detalhada — {len(resultados)} candidato(s)")
    ranking = st.session_state.get("ranking_base", ranking_base)
    nomes_ordenados = [(nome, veredicto) for nome, _, veredicto in ranking]
    abas = st.tabs([
        f"{medalha(i+1)} {'✅' if v == 'Aprovado' else '❌'} {n.rsplit('.', 1)[0]}"
        for i, (n, v) in enumerate(nomes_ordenados)
    ])

    for aba, (nome, veredicto) in zip(abas, nomes_ordenados):
        with aba:
            dados = resultados[nome]
            sub_analise, sub_perguntas, sub_roteiro = st.tabs([
                "📊 Análise",
                "❓ Perguntas de Entrevista",
                "📋 Roteiro de Entrevista",
            ])

            with sub_analise:
                if veredicto == "Aprovado":
                    st.success("✅ APROVADO — Candidato atende aos critérios da vaga")
                elif veredicto == "Reprovado":
                    st.error("❌ REPROVADO — Candidato não atende aos critérios da vaga")
                st.markdown(dados["analise"])
                st.download_button(
                    label="⬇️ Baixar Análise (TXT)",
                    data=dados["analise"],
                    file_name=f"analise_{nome.rsplit('.', 1)[0]}.txt",
                    mime="text/plain",
                    key=f"dl_analise_{nome}",
                )

            with sub_perguntas:
                st.markdown(dados["perguntas"])
                st.download_button(
                    label="⬇️ Baixar Perguntas (TXT)",
                    data=dados["perguntas"],
                    file_name=f"perguntas_{nome.rsplit('.', 1)[0]}.txt",
                    mime="text/plain",
                    key=f"dl_perguntas_{nome}",
                )

            with sub_roteiro:
                chave_roteiro = f"roteiro_{nome}"
                if chave_roteiro in st.session_state:
                    st.markdown(st.session_state[chave_roteiro])
                    st.download_button(
                        label="⬇️ Baixar Roteiro (TXT)",
                        data=st.session_state[chave_roteiro],
                        file_name=f"roteiro_{nome.rsplit('.', 1)[0]}.txt",
                        mime="text/plain",
                        key=f"dl_roteiro_{nome}",
                    )
                else:
                    st.info(
                        "O roteiro é gerado sob demanda — clique abaixo para criar o roteiro "
                        "completo da entrevista com este candidato."
                    )
                    if st.button(
                        "📋 Gerar Roteiro de Entrevista",
                        key=f"btn_roteiro_{nome}",
                        type="primary",
                        use_container_width=True,
                    ):
                        texto_cv = dados.get("texto_cv", "Currículo não disponível.")
                        with st.spinner("Gerando roteiro personalizado..."):
                            roteiro = gerar_roteiro_entrevista(
                                texto_cv,
                                dados["analise"],
                                descricao_vaga,
                                expectativas,
                            )
                        st.session_state[chave_roteiro] = roteiro
                        st.rerun()
